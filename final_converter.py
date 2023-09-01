import re
import sys
from docx import Document
#this script converts a RIS file exported from an EndNote library into unformatted Word CWYW citations
doi_pat = re.compile(r'\b10\.\d{4,}(?:\.\d+)*\/[-._;()/:A-Za-z0-9]+\b')

try:
    filename = input("Enter the name of the input RIS file exported from the EndNote library linked to the Word Document: ")
    content = ''
    try:
        with open(filename, encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError as e:
        print("Error decoding the file:", e)
        sys.exit(1)
except Exception as e:
    print("An error occurred while opening the file:", e)
    sys.exit(1)

refs = content.split('ER  - ')#er is not included in theproduced string 

#dictionary that stores the doi as the key and the unformatted citation as the value
conv_refs = {}


for ref in refs:
  #assuming there is just one DOI per endnote reference
    if (ref.strip() == ''):
        continue
    match = doi_pat.search(ref)
    if match:
      doi =  (match.group(0)).lower()
      
      conv = '{'
      #extract author last name (if present)
      i1 = ref.find('AU  -')
      i2 = ref.find(', ', i1 + 6)
      if i1 != -1 and i2 != -1:
        conv += (ref[i1+6:i2] + ', ')
        
      #extract publication year (if present)
      i1 = ref.find('PY  -')
      i2 = ref.find('\n', i1 + 6)
      if (i1 != -1 and i2 != -1):
        conv += (ref[i1+6:i2] + ' ')
        
      #extract ref number
      i1 = ref.find('ID  -')
      i2 = ref.find('\n', i1 + 6)
      if (i1 != -1 and i2 != -1):
        conv += ( '#' +  ref[i1+6:i2] + '}')
        conv_refs[doi] = conv

'''
#writing the dictionary values to an output file for testing 
with open('output.txt', 'w', encoding='utf-8') as f:
    for key, value in conv_refs.items():
        f.write(key + ' : ' + value + '\n')
'''

#------------------word doc stuff-------------------------------
word_fname = input("Enter the name of the input Word Document: ")
try:
    doc = Document(word_fname)
except Exception as e:
    print("An error occurred while initializing the Document object:", e)
    sys.exit(1)

#iterate through paragraphs
for paragraph in doc.paragraphs:
    dois = doi_pat.findall(paragraph.text)
    for doi in dois:
        if doi.lower() in conv_refs:
           paragraph.text = paragraph.text.replace(doi, conv_refs[doi.lower()])
        else:
            with open('missing_refs.txt', 'a') as f:
              f.write(doi + '\n')
  

# Iterate through tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            dois = doi_pat.findall(cell.text)
            for doi in dois:
                if doi.lower() in conv_refs:
                    cell.text = cell.text.replace(doi, conv_refs[doi.lower()])
                else:
                    with open('missing_refs.txt', 'a') as f:
                        f.write(doi + '\n')
                        
doc.save('new_' + word_fname)
print('All done, the new file has been saved as new_' + word_fname)
print('unmatched DOIs have been written to the file missing_refs.txt')
